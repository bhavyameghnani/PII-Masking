import re
import json
import boto3
import os
from docx import Document as DocxDocument
from langchain_experimental.data_anonymizer import PresidioReversibleAnonymizer
from presidio_analyzer import Pattern, PatternRecognizer
from faker import Faker
from presidio_anonymizer.entities import OperatorConfig
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.schema import Document
from operator import itemgetter
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import (
    RunnableLambda,
    RunnableParallel,
    RunnablePassthrough,
)
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()
AWS_REGION = os.getenv("AWS_REGION")
BUCKET_NAME = os.getenv("BUCKET_NAME")
DOCX_KEY = os.getenv("DOCX_KEY")  # The S3 key for the DOCX file
EMBEDDINGS_KEY = os.getenv("EMBEDDINGS_KEY", "embeddings.faiss")  # S3 key for embeddings
ANONYMIZATION_MAP_KEY = os.getenv("ANONYMIZATION_MAP_KEY", "anonymization_map.json")  # S3 key for the anonymization map

# Initialize AWS clients
s3_client = boto3.client("s3", region_name=AWS_REGION)
bedrock_client = boto3.client(service_name="bedrock-runtime", region_name=AWS_REGION)
bedrock_embeddings = BedrockEmbeddings(model_id="amazon.titan-embed-text-v1", client=bedrock_client)

# Function to read .docx file and return the text content
def read_docx(file_path):
    doc = DocxDocument(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

# Download the DOCX file from S3
s3_client.download_file(BUCKET_NAME, DOCX_KEY, '/tmp/document.docx')

# Load and read the DOCX file
document_content = read_docx('/tmp/document.docx')

# Define patterns for Polish ID and time
polish_id_pattern = Pattern(
    name="polish_id_pattern",
    regex="[A-Z]{3}\\d{6}",
    score=1,
)
time_pattern = Pattern(
    name="time_pattern",
    regex="(1[0-2]|0?[1-9]):[0-5][0-9] (AM|PM)",
    score=1,
)

# Define the recognizers with the patterns
polish_id_recognizer = PatternRecognizer(
    supported_entity="POLISH_ID", patterns=[polish_id_pattern]
)
time_recognizer = PatternRecognizer(supported_entity="TIME", patterns=[time_pattern])

# Initialize Faker for custom fake data generation
fake = Faker()

# Custom function to generate fake Polish ID
def fake_polish_id(_=None):
    return fake.bothify(text="???######").upper()

# Test the fake Polish ID function
fake_polish_id()

# Custom function to generate fake time
def fake_time(_=None):
    return fake.time(pattern="%I:%M %p")

# Test the fake time function
fake_time()

# Define custom operators for the anonymizer
new_operators = {
    "POLISH_ID": OperatorConfig("custom", {"lambda": fake_polish_id}),
    "TIME": OperatorConfig("custom", {"lambda": fake_time}),
}

# Initialize the anonymizer again with a seed for reproducibility
anonymizer = PresidioReversibleAnonymizer(
    faker_seed=42,
)

# Add the custom recognizers and operators again
anonymizer.add_recognizer(polish_id_recognizer)
anonymizer.add_recognizer(time_recognizer)
anonymizer.add_operators(new_operators)

# Anonymize the document before indexing
anonymized_content = anonymizer.anonymize(document_content)

# Extract the anonymization map to store in JSON
anonymization_map = anonymizer.deanonymizer_mapping

# Save the anonymization map to a JSON file and upload to S3
with open('/tmp/anonymization_map.json', 'w') as f:
    json.dump(anonymization_map, f, indent=4)
s3_client.upload_file('/tmp/anonymization_map.json', BUCKET_NAME, ANONYMIZATION_MAP_KEY)

# Split the anonymized content into chunks
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
chunks = text_splitter.split_text(anonymized_content)

# Convert chunks to Document objects
documents = [Document(page_content=chunk) for chunk in chunks]

# Index the chunks using Bedrock embeddings
docsearch = FAISS.from_documents(documents, bedrock_embeddings)

# Save the FAISS index to disk and upload to S3
faiss_index_path = '/tmp/embeddings.faiss'
docsearch.save_local(faiss_index_path)
s3_client.upload_file(faiss_index_path, BUCKET_NAME, EMBEDDINGS_KEY)

# Later, when you need to query the stored embeddings

# Download the FAISS index and anonymization map from S3
s3_client.download_file(BUCKET_NAME, EMBEDDINGS_KEY, faiss_index_path)
s3_client.download_file(BUCKET_NAME, ANONYMIZATION_MAP_KEY, '/tmp/anonymization_map.json')

# Load the FAISS index
retrieved_docsearch = FAISS.load_local(faiss_index_path, bedrock_embeddings)

# Load the anonymization map
with open('/tmp/anonymization_map.json', 'r') as f:
    anonymization_map = json.load(f)

# Initialize the anonymizer with the loaded anonymization map
anonymizer = PresidioReversibleAnonymizer(
    faker_seed=42,
    deanonymizer_mapping=anonymization_map
)

# Create the retriever
retriever = retrieved_docsearch.as_retriever()

# Create an anonymizer chain with prompt template and Bedrock model
template = """Answer the question based only on the following context:
{context}

Question: {anonymized_question}
"""
prompt = ChatPromptTemplate.from_template(template)

model = BedrockChat(model_id="anthropic.claude-3-5-sonnet-20240620-v1:0", client=bedrock_client)

# Define parallel input processing for the chain
_inputs = RunnableParallel(
    question=RunnablePassthrough(),
    anonymized_question=RunnableLambda(anonymizer.anonymize),
)

# Create the anonymizer chain
anonymizer_chain = (
    _inputs
    | {
        "context": itemgetter("anonymized_question") | retriever,
        "anonymized_question": itemgetter("anonymized_question"),
    }
    | prompt
    | model
    | StrOutputParser()
)

# Add deanonymization step to the chain
chain_with_deanonymization = anonymizer_chain | RunnableLambda(anonymizer.deanonymize)

# Invoke the chain with deanonymization and print the results
print(
    chain_with_deanonymization.invoke(
        "Which Company is a Party A and which company is Party B in this agreement?"
    )
)

print(
    chain_with_deanonymization.invoke("List all “Specified Entity” means in relation to Party A for the purpose of Section 5(a)(v),")
)

print(chain_with_deanonymization.invoke("Please summarise Credit Event Upon Merger clause"))
